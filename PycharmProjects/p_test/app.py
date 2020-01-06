# coding=UTF-8
from flask import Flask, render_template, jsonify, request
import os, time
from pyecharts.charts import Grid, Bar, Page
from pyecharts import options as opts
from scripts import word_functions
from docx import Document
from pyecharts.globals import ThemeType
app = Flask(__name__)
result_list = []
state = 'no_ready'
file = None
file1 = None
result_num_list = []
ip = None
date = None
select_test = None


@app.route('/')
@app.route('/main/')
def first():
    return render_template('main.html')

@app.route('/wait/')
def wait():
    return render_template('wait.html')


@app.route('/result_num/', methods=['GET'])
def result_num():
    return jsonify({'code': 200, 'message': '创建报告成功'})


@app.route('/result/')
def result():
    #time.sleep(0.5)
    global file, state, result_list, ip
    while True:
        if os.path.exists(file):
            break
        else:
            continue
    with open(file, "r") as f:
        result_list = f.readlines()
    return render_template('result.html', result=result_list, state=state, ip=ip, select_test=select_test)


@app.route('/start/', methods=['POST'])
def start():
    print("---------------------")
    global file, state, file1, ip, date ,select_test
    date = time.strftime("%Y%m%d%H%M%S", time.localtime())
    file = "/root/PycharmProjects/p_test/logs/" + date + ".txt"
    file1 = "/root/PycharmProjects/p_test/logs/" + date + "+.txt"
    idd = request.form.get('idd')
    disk = request.form.get('disk')
    ip_1 = request.form.get('ip_1')
    ip_2 = request.form.get('ip_2')

    ip = os.popen("ifconfig wlp3s0| awk '/inet / {print $2}' | cut -f2 -d:").read().strip('\n')

    if idd == 'start_net':
        os.popen('bash /root/PycharmProjects/p_test/net_test.sh ' + date + ' ' + ip_1 + ' ' + ip_2 + ' ' + disk).read()

    if idd == 'start_r4k':
        os.popen('bash /root/PycharmProjects/p_test/r4k_test.sh ' + date + ' ' + ip_1 + ' ' + ip_2 + ' ' + disk).read()

    if idd == 'start_w4k':
        os.popen('bash /root/PycharmProjects/p_test/w4k_test.sh ' + date + ' ' + ip_1 + ' ' + ip_2 + ' ' + disk).read()

    if idd == 'start_r64k':
        os.popen('bash /root/PycharmProjects/p_test/r64k_test.sh ' + date + ' ' + ip_1 + ' ' + ip_2 + ' ' + disk).read()

    if idd == 'start_w64k':
        os.popen('bash /root/PycharmProjects/p_test/w64k_test.sh ' + date + ' ' + ip_1 + ' ' + ip_2 + ' ' + disk).read()

    if idd == 'start_rw64k':
        os.popen('bash /root/PycharmProjects/p_test/rw64k_test.sh ' + date + ' ' + ip_1 + ' ' + ip_2 + ' ' + disk).read()

    if idd == 'start_r1m':
        os.popen('bash /root/PycharmProjects/p_test/r1m_test.sh ' + date + ' ' + ip_1 + ' ' + ip_2 + ' ' + disk).read()

    if idd == 'start_w1m':
        os.popen('bash /root/PycharmProjects/p_test/w1m_test.sh ' + date + ' ' + ip_1 + ' ' + ip_2 + ' ' + disk).read()

    if idd == 'start_rw1m':
        os.popen('bash /root/PycharmProjects/p_test/rw1m_test.sh ' + date + ' ' + ip_1 + ' ' + ip_2 + ' ' + disk).read()

    if idd == 'all':
        os.popen('bash /root/PycharmProjects/p_test/all_test.sh ' + date + ' ' + ip_1 + ' ' + ip_2 + ' ' + disk).read()

        def unify(list):
            if list[-3:] == 'B/s' and list[-4] != 'K' and list[-4] != 'G' and list[-4] != 'M':
                date = int(float(list[0:-3]) / (1024 * 1024))
            else:
                unit = list[-4:]
                date = list[0:-4]
                if unit == 'MB/s':
                    date = int(float(date))
                elif unit == 'KB/s':
                    date = int(float(date) / 1024)
                elif unit == 'GB/s':
                    date = int(float(date)) * 1024
                elif unit == '/sec':
                    date = list[0:-9]
                    date = int(float(date))
                else:
                    date = int(list)
            return date
        while True:
            if os.path.exists(file1):
                break
            else:
                continue
        with open(file1, "r") as f:
            for line in f.readlines():
                line = line.strip('\n')
                a = unify(line)
                result_num_list.append(a)
        print(result_num_list)

        bar1 = (
            Bar(init_opts=opts.InitOpts(theme=ThemeType.PURPLE_PASSION))
                .add_xaxis(["ip1-ip2", "ip2-ip1"])
                .add_yaxis("", [result_num_list[0], result_num_list[1]], category_gap="80%")
                .set_global_opts(title_opts=opts.TitleOpts(title="Network Bandwidth Performance", subtitle="Mbits/sec"))
        )

        bar2 = (
            Bar(init_opts=opts.InitOpts(theme=ThemeType.PURPLE_PASSION))
                .add_xaxis(["4K", "64K", "1M"])

                .add_yaxis("Read", [max(result_num_list[2], result_num_list[4], result_num_list[6]),
                                    max(result_num_list[26], result_num_list[28], result_num_list[30]),
                                    max(result_num_list[50], result_num_list[52], result_num_list[54])],
                           category_gap="80%")
                .set_global_opts(title_opts=opts.TitleOpts(title="IOPS - 100% Read"),
                                 legend_opts=opts.LegendOpts(pos_left="50%"))
        )

        bar3 = Bar(init_opts=opts.InitOpts(theme=ThemeType.PURPLE_PASSION))
        bar3.add_xaxis(["4K", "64K", "1M"])
        bar3.add_yaxis("Read", [max(result_num_list[3], result_num_list[5], result_num_list[7]),
                                max(result_num_list[27], result_num_list[29], result_num_list[31]),
                                max(result_num_list[51], result_num_list[53], result_num_list[55])],
                       category_gap="80%")
        bar3.set_global_opts(title_opts=opts.TitleOpts(title="IOBW - 100% Read", subtitle="MB/s"),
                             legend_opts=opts.LegendOpts(pos_left="50%"))

        bar4 = Bar(init_opts=opts.InitOpts(theme=ThemeType.PURPLE_PASSION))
        bar4.add_xaxis(["4K", "64K", "1M"])
        bar4.add_yaxis("Write", [max(result_num_list[8], result_num_list[10], result_num_list[12]),
                                 max(result_num_list[32], result_num_list[34], result_num_list[36]),
                                 max(result_num_list[56], result_num_list[58], result_num_list[60])],
                       category_gap="80%")
        bar4.set_global_opts(title_opts=opts.TitleOpts(title="IOPS - 100% Write"),
                             legend_opts=opts.LegendOpts(pos_left="50%"))

        bar5 = Bar(init_opts=opts.InitOpts(theme=ThemeType.PURPLE_PASSION))
        bar5.add_xaxis(["4K", "64K", "1M"])
        bar5.add_yaxis("Write", [max(result_num_list[9], result_num_list[11], result_num_list[13]),
                                 max(result_num_list[33], result_num_list[35], result_num_list[37]),
                                 max(result_num_list[57], result_num_list[59], result_num_list[61])],
                       category_gap="80%")
        bar5.set_global_opts(title_opts=opts.TitleOpts(title="IOBW - 100% Write", subtitle="MB/s"),
                             legend_opts=opts.LegendOpts(pos_left="50%"))

        bar6 = Bar(init_opts=opts.InitOpts(theme=ThemeType.PURPLE_PASSION))
        bar6.add_xaxis(["4K", "64K", "1M"])
        bar6.add_yaxis("Write", [max(result_num_list[15], result_num_list[19], result_num_list[23]),
                                 max(result_num_list[39], result_num_list[43], result_num_list[47]),
                                 max(result_num_list[63], result_num_list[67], result_num_list[71])],
                       category_gap="80%")
        bar6.add_yaxis("Read", [max(result_num_list[14], result_num_list[18], result_num_list[22]),
                                max(result_num_list[38], result_num_list[42], result_num_list[46]),
                                max(result_num_list[62], result_num_list[66], result_num_list[70])],
                       category_gap="80%")

        bar6.set_global_opts(title_opts=opts.TitleOpts(title="IOPS - 50% Read 50% Write"),
                             legend_opts=opts.LegendOpts(pos_left="50%"))

        bar7 = Bar(init_opts=opts.InitOpts(theme=ThemeType.PURPLE_PASSION))
        bar7.add_xaxis(["4K", "64K", "1M"])
        bar7.add_yaxis("Write", [max(result_num_list[17], result_num_list[21], result_num_list[25]),
                                 max(result_num_list[41], result_num_list[45], result_num_list[49]),
                                 max(result_num_list[65], result_num_list[69], result_num_list[73])],
                       category_gap="80%")
        bar7.add_yaxis("Read", [max(result_num_list[16], result_num_list[20], result_num_list[24]),
                                max(result_num_list[40], result_num_list[44], result_num_list[48]),
                                max(result_num_list[64], result_num_list[68], result_num_list[72])],
                       category_gap="80%")

        bar7.set_global_opts(title_opts=opts.TitleOpts(title="IOBW - 50% Read 50% Write", subtitle="MB/s"),
                             legend_opts=opts.LegendOpts(pos_left="50%"))
        os.popen('mkdir /var/www/html/results/' + date)
        bar1.render(path='/var/www/html/results/' + date + '/bar1.html')
        bar2.render(path='/var/www/html/results/' + date + '/bar2.html')
        bar3.render(path='/var/www/html/results/' + date + '/bar3.html')
        bar4.render(path='/var/www/html/results/' + date + '/bar4.html')
        bar5.render(path='/var/www/html/results/' + date + '/bar5.html')
        bar6.render(path='/var/www/html/results/' + date + '/bar6.html')
        bar7.render(path='/var/www/html/results/' + date + '/bar7.html')
        time.sleep(5)
        for i in [1, 2, 3, 4, 5, 6, 7]:
            print(i)
            cmd = "sed -i 's/https\:\/\/assets\.pyecharts\.org\/assets/\/js/' /var/www/html/results/" + date + "/bar" + str(
                i) + ".html"
            os.popen(cmd)
            cmd1 = "webscreenshot http://127.0.0.1/results/" + date + "/bar" + str(
                i) + ".html -o /var/www/html/results/" + date + "/png --window-size '800,500'"
            os.popen(cmd1).read()
            cmd2 = "mv /var/www/html/results/" + date + "/png/*bar" + str(
                i) + "* /var/www/html/results/" + date + "/png/bar" + str(i) + ".png"
            print(cmd2)
            os.popen(cmd2).read()
        a = Document("./poc.docx")
        a = word_functions.docx_replace_table_text("网络性能测试结果", "__network_test_up_bw__", str(result_num_list[0]), a)
        a = word_functions.docx_replace_table_text("网络性能测试结果", "__network_test_down_bw__", str(result_num_list[1]), a)

        a = word_functions.docx_replace_table_text("存储性能测试结果", "__4k_rand_100r_bw__",
                                                   str(max(result_num_list[3], result_num_list[5], result_num_list[7])),
                                                   a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__4k_rand_100r_iops__",
                                                   str(max(result_num_list[2], result_num_list[4], result_num_list[6])),
                                                   a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__4k_rand_100w_bw__",
                                                   str(max(result_num_list[9], result_num_list[11],
                                                           result_num_list[13])),
                                                   a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__4k_rand_100w_iops__",
                                                   str(max(result_num_list[8], result_num_list[10],
                                                           result_num_list[12])),
                                                   a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__4k_rand_50r50w_bw__", '读:' + str(
            max(result_num_list[16], result_num_list[20], result_num_list[24])) + ",写:" + str(
            max(result_num_list[17], result_num_list[21], result_num_list[25])), a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__4k_rand_50r50w_iops__", '读:' + str(
            max(result_num_list[14], result_num_list[18], result_num_list[22])) + ",写:" + str(
            max(result_num_list[15], result_num_list[19], result_num_list[23])), a)

        a = word_functions.docx_replace_table_text("存储性能测试结果", "__64k_rand_100r_bw__",
                                                   str(max(result_num_list[27], result_num_list[29],
                                                           result_num_list[31])),
                                                   a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__64k_rand_100r_iops__",
                                                   str(max(result_num_list[26], result_num_list[28],
                                                           result_num_list[30])),
                                                   a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__64k_rand_100w_bw__",
                                                   str(max(result_num_list[33], result_num_list[35],
                                                           result_num_list[37])),
                                                   a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__64k_rand_100w_iops__",
                                                   str(max(result_num_list[32], result_num_list[34],
                                                           result_num_list[36])),
                                                   a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__64k_rand_50r50w_bw__", '读:' + str(
            max(result_num_list[40], result_num_list[44], result_num_list[46])) + ",写:" + str(
            max(result_num_list[41], result_num_list[45], result_num_list[49])), a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__64k_rand_50r50w_iops__", '读:' + str(
            max(result_num_list[38], result_num_list[42], result_num_list[46])) + ",写:" + str(
            max(result_num_list[39], result_num_list[43], result_num_list[47])), a)

        a = word_functions.docx_replace_table_text("存储性能测试结果", "__1m_rand_100r_bw__",
                                                   str(max(result_num_list[51], result_num_list[53],
                                                           result_num_list[55])),
                                                   a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__1m_rand_100r_iops__",
                                                   str(max(result_num_list[50], result_num_list[52],
                                                           result_num_list[54])),
                                                   a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__1m_rand_100w_bw__",
                                                   str(max(result_num_list[57], result_num_list[59],
                                                           result_num_list[61])),
                                                   a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__1m_rand_100w_iops__",
                                                   str(max(result_num_list[56], result_num_list[58],
                                                           result_num_list[60])),
                                                   a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__1m_rand_50r50w_bw__", '读:' + str(
            max(result_num_list[64], result_num_list[68], result_num_list[72])) + ",写:" + str(
            max(result_num_list[65], result_num_list[69], result_num_list[73])), a)
        a = word_functions.docx_replace_table_text("存储性能测试结果", "__1m_rand_50r50w_iops__", '读:' + str(
            max(result_num_list[62], result_num_list[66], result_num_list[70])) + ",写:" + str(
            max(result_num_list[63], result_num_list[67], result_num_list[71])), a)

        a = word_functions.docx_add_picture("__network_test_pic_1__", "/var/www/html/results/" + date + "/png/bar1.png",
                                            a)

        a = word_functions.docx_add_picture("__io_test_pic_1__", "/var/www/html/results/" + date + "/png/bar2.png", a)
        a = word_functions.docx_add_picture("__io_test_pic_2__", "/var/www/html/results/" + date + "/png/bar3.png", a)
        a = word_functions.docx_add_picture("__io_test_pic_3__", "/var/www/html/results/" + date + "/png/bar4.png", a)
        a = word_functions.docx_add_picture("__io_test_pic_4__", "/var/www/html/results/" + date + "/png/bar5.png", a)
        a = word_functions.docx_add_picture("__io_test_pic_5__", "/var/www/html/results/" + date + "/png/bar6.png", a)
        a = word_functions.docx_add_picture("__io_test_pic_6__", "/var/www/html/results/" + date + "/png/bar7.png", a)

        a.save("/var/www/html/results/" + date + "/POC检测报告.docx")


        select_test = 'all'

    state = 'ready'
    return jsonify({'code': 200, 'message': '234'})


@app.route('/restart/', methods=['POST'])
def restart():
    global result_list, state, file, file1, result_num_list, ip
    ip = None
    result_num_list = []
    file1 = None
    result_list = []
    state = 'no_ready'
    file = None
    date = None
    return jsonify({'code': 200, 'message': '234'})

@app.route('/clear/')
def clear():
    global result_list, state, file, file1, result_num_list, ip
    cmd = "rm -rf ./logs/* && rm -rf /var/www/html/results/*"
    os.popen(cmd)
    return jsonify({'code': 200, 'message': '234'})


if __name__ == '__main__':
    app.run(host='0.0.0.0')