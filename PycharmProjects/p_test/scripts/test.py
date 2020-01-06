#!/usr/local/bin/python3
# encoding: UTF-8

import uuid
import word_functions
from datetime import datetime
from docxtpl import DocxTemplate
from shutil import copyfile
from docx import Document
import os 

# 输入参数
hybrid_intro = "yes"
disaster_recover_live_migration = "yes"
project_architecture = "ceph"
rack_architecture = "ceph"
var_list = {'company_name': '上海融计', 'project_date': datetime.now().strftime('%Y-%m-%d'), }

# 准备工作空间
tmp_file = "/tmp/" + uuid.uuid1().hex.upper()[0:6] + ".docx"
copyfile("../ZStack-template.docx", tmp_file)
doc_tmp_file = Document(tmp_file)

# Phase 1: 拼接文件
if hybrid_intro == "yes":
    doc = Document("../hybrid_intro.docx")
    doc_tmp_file = word_functions.docx_insert_word_files(doc_tmp_file, doc, "__hybrid_intro__")

if disaster_recover_live_migration == "yes":
    doc = Document("../disaster_recover_live_migration.docx")
    doc_tmp_file = word_functions.docx_insert_word_files(doc_tmp_file, doc, "__disaster_recover_live_migration__")

if project_architecture == "ceph":
    doc = Document("../ceph_arch/ceph_arch_explain.docx")
    doc_tmp_file = word_functions.docx_insert_word_files(doc_tmp_file, doc, "__arch_explain__")

if hybrid_intro == "yes":
    doc = Document("../hybrid_feature_list.docx")
    doc_tmp_file = word_functions.docx_insert_word_files(doc_tmp_file, doc, "__feature_list__")


# Phase 2: 替换图片
doc_tmp_file = word_functions.docx_add_picture("project_architecture", "../ceph_arch/ceph_arch.jpg", doc_tmp_file)

doc_tmp_file = word_functions.docx_add_picture("rack_architecture", "../ceph_arch/ceph_arch.jpg", doc_tmp_file)

doc_tmp_file.save(tmp_file)

# Phase 3: 替换文字
word_functions.docxtpl_render_text(DocxTemplate(tmp_file), var_list).save(tmp_file)
copyfile(tmp_file, "../test.docx")
os.remove(tmp_file)
