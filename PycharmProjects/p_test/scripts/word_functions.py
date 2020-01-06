#!/usr/local/bin/python3
# encoding: UTF-8

# Reference: https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html

# docxtpl外输入文件输出文件外，docx输入doc输出doc

import re
import sys
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import qn


# Phase 1
# Base functions
def docx_replace_text(var_dest, content, doc):
    try:
        for p in doc.paragraphs:
            if re.search(var_dest, p.text):
                p.text = re.sub(var_dest, ("{0}").format(content), p.text)
    except Exception as e:
        print(e)
        return 1
    return doc

def delete_table(var_table_title, doc):
    try:
        for i in range(0, len(doc.tables)):
            if doc.tables[i].row_cells(0)[0].text == var_table_title:
                doc.tables.pop(i)
                print("Delete table {0}"+var_table_title)
    except Exception as e:
        print(e)
        return 1
    return doc

# 还是用docxtpl吧
def docx_replace_table_text(var_table_title, var_dest, content, doc):
    try:
        for t in doc.tables:
            if t.row_cells(0)[0].text == var_table_title:
                for i in range(0, len(t.rows)):
                    for j in range(0, len(t.row_cells(i))):
                        if re.search(var_dest, t.row_cells(i)[j].text):
                            t.row_cells(i)[j].text = re.sub(var_dest, ("{0}").format(content), t.row_cells(i)[j].text)
    except Exception as e:
        print(e)
        return 1
    return doc

def docx_add_picture(var_dest, pic_path, doc):
    try:
        for p in doc.paragraphs:
            if re.search(var_dest, p.text):
                p.clear()
                r = p.add_run()
                r.add_picture(pic_path, width=Inches(6))
                p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        print(e)
    return doc

def docx_combine_word_files(files, dest):
    #files = ['file1.docx', 'file2.docx']
    merged_document = Document()
    for index, file in enumerate(files):
        sub_doc = file
        # Don't add a page break if you've reached the last file.
        if index < len(files)-1:
           sub_doc.add_page_break()
        for element in sub_doc.element.body:
            merged_document.element.body.append(element)
    return merged_document

def docx_insert_word_files(main_doc, sub_doc, var_dest):
    try:
        for p in main_doc.element.body.findall('.//'+qn('w:p')):
            if re.search(var_dest, ''.join([t.text for t in p.findall('.//'+qn('w:t'))])):
                p.clear()
                print("Insert word file at: {0} ".format(var_dest))
                for element in sub_doc.element.body:
                    p.addprevious(element)
    except Exception as e:
        print(e)
    return main_doc

def docx_add_table():
    try:
        t=f.add_table(3,3)
        t.style='Table Grid'
    except Exception as e:
        print(e)
    return 0

# Phase 2
def docxtpl_render_text(doc, var_list):
    context = var_list
    doc.render(context)
    return doc
